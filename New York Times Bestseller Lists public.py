'''New York Times Bestseller Lists.py
This project reduces the amount of copying and pasting NYT Bestseller
data from the website to a document template.
'''


import requests, bs4
from docx import Document
from docx.shared import Pt
#download website data
result_fic, result_nonfic, result_advice = requests.get("https://www.nytimes.com/books/best-sellers/hardcover-fiction/"),requests.get("https://www.nytimes.com/books/best-sellers/hardcover-nonfiction/"),requests.get("https://www.nytimes.com/books/best-sellers/advice-how-to-and-miscellaneous/")
#check if downloaded correctly
result_fic.raise_for_status()
result_nonfic.raise_for_status()
result_advice.raise_for_status()
#turn content into BS4 objects
c_fic = result_fic.content
soup_fic = bs4.BeautifulSoup(c_fic, features="lxml")

c_nonfic = result_nonfic.content
soup_nonfic = bs4.BeautifulSoup(c_nonfic, features="lxml")

c_advice = result_advice.content
soup_advice = bs4.BeautifulSoup(c_advice, features="lxml")

#empty lists
publisher_text = []
title_text = []
author_text = []
desc_text = []
fresh_text = []

#open doc
document = Document('NYT Bestsellers blank and placeholder logo.docx') 

def assign_tags(soup):
    publisher_tag = soup.select('.publisher')
    title_tag = soup.find_all("h2", "title") #select does not work
    author_tag = soup.select('.author')
    desc_tag = soup.select('.description')
    fresh_tag = soup.select('.freshness')

    for i in range(len(publisher_tag)):
        publisher_text.append(publisher_tag[i].text)
        title_text.append(title_tag[i].text)
        author_text.append(author_tag[i].text)
        desc_text.append(desc_tag[i].text.strip(" "))
        fresh_text.append((fresh_tag[i].text.upper())[0:3])
        for j in "W ": #we don't need the extra "W" and " " that come with the slice
            if fresh_text[i] == "NEW": #but we don't want to accidentally mutate "NEW"
                continue
            else:
                fresh_text[i] = fresh_text[i].replace(j, '')
    return fresh_text, author_text, desc_text, title_text, publisher_text
        


#looping to list streamlining--function or just do it?
def make_doc(fresh_text, author_text, desc_text, title_text, publisher_text, table_no):
    table = document.tables[table_no] #establish table number
    print(table_no)

    row_count = len(table.rows)
    for item in range(0, row_count-1):
        row = table.rows[item+1]
        p = row.cells[3].paragraphs[0] #find first paragraph (all cells have at least one)
        #paragraph formatting
        p.style = "Normal"
        p.paragraph_format.space_after = Pt(0)
        #first runner content and formatting
        title = title_text[item]
        runner = p.add_run(title)
        runner.bold = True
        runner.underline = True
        runner.font.name = 'Times New Roman'
        runner.font.size = Pt(11)
        #second runner content and formatting
        cell_content = p.add_run(" " + author_text[item] + " (" + publisher_text[item] + ") " + desc_text[item])
        cell_content.font.name = 'Times New Roman'
        cell_content.font.size = Pt(11)
        #freshness column content and formatting
        p2 = row.cells[4].paragraphs[0]
        fresh_content = p2.add_run(fresh_text[item])
        fresh_content.font.name = 'Times New Roman'
        document.save('NYT Bestsellers.docx')
    title_text.clear()
    fresh_text.clear()
    author_text.clear()
    desc_text.clear()
    publisher_text.clear()


assign_tags(soup_fic)
make_doc(fresh_text, author_text, desc_text, title_text, publisher_text, 0)
assign_tags(soup_nonfic)
make_doc(fresh_text, author_text, desc_text, title_text, publisher_text, 1)
assign_tags(soup_advice)
make_doc(fresh_text, author_text, desc_text, title_text, publisher_text, 2)
