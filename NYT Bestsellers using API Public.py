import requests, json, time
from docx import Document
from docx.shared import Pt
key = 'YOUR KEY HERE'
#open document
document = Document('NYT Bestsellers blank and placeholder logo.docx')
#empty lists
author_text = []
desc_text = []
publisher_text = []
weeks_text = []
title_text = []

def get_response_data(listname):
    #Get data from the bestseller list
    response = requests.get("https://api.nytimes.com/svc/books/v3/lists.json?list=" + listname + "&api-key=" + key)
    response.raise_for_status()
    response_data = response.json()
    time.sleep(1)
    return response_data

def parse_data(data):
    #Put necessary data into lists
    list_data = data
    for i in range(list_data['num_results']):
        global author_text, desc_text, publisher_text, weeks_text, title_text #not the best but it works
        author_text.append(list_data['results'][i]['book_details'][0]['contributor'])
        desc_text.append(list_data['results'][i]['book_details'][0]['description'])
        publisher_text.append(list_data['results'][i]['book_details'][0]['publisher'])
        print(weeks_text)
        if list_data['results'][i]['weeks_on_list'] == 1:
            weeks_text.append("NEW")
        else:
            weeks_text.append(list_data['results'][i]['weeks_on_list'])
        title_text.append(list_data['results'][i]['book_details'][0]['title'])
    return author_text, desc_text, publisher_text, weeks_text, title_text

def make_doc(author_text, desc_text, publisher_text, weeks_text, title_text, table_no):
    
    table = document.tables[table_no] #establish table number
    print(table_no)
    print(title_text)

    row_count = len(table.rows)
    for item in range(0, row_count-1):
        row = table.rows[item+1]
        p = row.cells[3].paragraphs[0] #all cells have at least one paragraph
        #paragraph formatting
        p.style = "Normal"
        p.paragraph_format.space_after = Pt(0)

        #first runner content and formatting
        title = title_text[item]
        runner = p.add_run(title)
        runner.bold = True
        runner.underline = True
        runner.font.name = 'Times New Roman'
        runner.font.size = Pt(11)
        #second runner content and formatting
        cell_content = p.add_run(" " + author_text[item] + " (" + publisher_text[item] + ") " + desc_text[item])
        cell_content.font.name = 'Times New Roman'
        cell_content.font.size = Pt(11)
        #freshness column content and formatting
        p2 = row.cells[4].paragraphs[0]
        weeks_content = p2.add_run(str(weeks_text[item])) #add_run works with strings only
        weeks_content.font.name = 'Times New Roman'
        document.save('NYT Bestsellers.docx')
    title_text.clear()
    author_text.clear()
    desc_text.clear()
    publisher_text.clear()
    weeks_text.clear()

def new_york_times():
    fiction = get_response_data('hardcover-fiction')
    parse_data(fiction)
    make_doc(author_text, desc_text, publisher_text, weeks_text, title_text, 0)
    
    nonfiction = get_response_data('hardcover-nonfiction')
    parse_data(nonfiction)
    make_doc(author_text, desc_text, publisher_text, weeks_text, title_text, 1)
    
    advice = get_response_data('advice-how-to-and-miscellaneous')
    parse_data(advice)
    make_doc(author_text, desc_text, publisher_text, weeks_text, title_text, 2)


new_york_times()
